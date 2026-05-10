#!/usr/bin/env python3
"""
generate_resume.py — Produces ATS-safe, human-quality .docx (and optional PDF)
resumes and cover letters from markdown content.

Usage:
    python3 scripts/generate_resume.py --company "Stripe" --content-file "$WORKSPACE/resume_content.md"
    python3 scripts/generate_resume.py --company "Stripe" --type cover-letter --content-file "$WORKSPACE/cover_letter_content.md"
    python3 scripts/generate_resume.py --company "Stripe" --content-file "..." --pdf   # also generate PDF

Output:
    $WORKSPACE/resume.docx
    $WORKSPACE/resume.pdf  (if --pdf)
"""

import argparse
import os
import re
import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ── Formatting constants (ATS-safe, A4, Irish market) ────────────────────────
FONT_NAME = "Calibri"
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_NAME = Pt(16)
FONT_SIZE_DESIGNATION = Pt(11)
FONT_SIZE_CONTACT = Pt(9.5)
FONT_SIZE_SECTION = Pt(10.5)
FONT_SIZE_ROLE = Pt(10)
FONT_SIZE_SUBROLE = Pt(10)
MARGIN = Mm(12.7)         # Word "Narrow" preset
MARGIN_TOP = Mm(12.7)
MARGIN_BOT = Mm(12.7)
LINE_SPACING_VAL = 276    # 1.15x line spacing in twips (auto rule)
BULLET_INDENT = 340       # twips — bullet left indent
SKILLS_INDENT = 3600      # twips — skills hanging indent

# A4 page dimensions
A4_WIDTH = Mm(210)
A4_HEIGHT = Mm(297)

# Colours — professional palette
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)        # Dark grey body text
NAME_COLOR = RGBColor(0x1F, 0x3A, 0x5F)        # Dark navy for name
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)     # Dark navy for section headings
ROLE_COLOR = RGBColor(0x2D, 0x2D, 0x2D)        # Near-black for role titles
DIVIDER_COLOR = "7F9BBF"                         # Muted blue-grey for section lines

# ── British English: US→UK replacements ──────────────────────────────────────
# Context-aware: "program" only replaced when NOT preceded by "computer/software/python/java/shell/c"
US_TO_UK = {
    "analyzed": "analysed", "analyzing": "analysing", "analyze": "analyse",
    "optimize": "optimise", "optimized": "optimised", "optimizing": "optimising",
    "optimization": "optimisation",
    "organize": "organise", "organized": "organised", "organizing": "organising",
    "organization": "organisation",
    "recognize": "recognise", "recognized": "recognised", "recognizing": "recognising",
    "utilize": "utilise", "utilized": "utilised", "utilizing": "utilising",
    "utilization": "utilisation",
    "customize": "customise", "customized": "customised", "customizing": "customising",
    "minimize": "minimise", "minimized": "minimised", "minimizing": "minimising",
    "maximize": "maximise", "maximized": "maximised", "maximizing": "maximising",
    "standardize": "standardise", "standardized": "standardised",
    "prioritize": "prioritise", "prioritized": "prioritised",
    "summarize": "summarise", "summarized": "summarised",
    "categorize": "categorise", "categorized": "categorised",
    "visualize": "visualise", "visualized": "visualised",
    "visualization": "visualisation",
    "color": "colour", "colored": "coloured",
    "favor": "favour", "favorable": "favourable",
    "behavior": "behaviour", "behavioral": "behavioural",
    "center": "centre", "centered": "centred",
    "license": "licence", "licensed": "licenced",
    "defense": "defence", "offense": "offence",
    "catalog": "catalogue",
    "dialog": "dialogue",
    "modeling": "modelling", "modeled": "modelled",
    "traveling": "travelling", "traveled": "travelled",
    "labeling": "labelling", "labeled": "labelled",
    "fulfillment": "fulfilment",
    "enrollment": "enrolment",
    "skillset": "skill set",
}

# "program" needs context-aware replacement — only replace when it means "scheme/initiative"
# NOT when it means "computer program" or "programming"
PROGRAM_SAFE_PREFIXES = re.compile(
    r'(?:computer|software|python|java|shell|c\+\+|coding|code)\s+program',
    re.IGNORECASE
)


def sanitise_text(text: str) -> str:
    """Remove invisible/ghost characters that trip ATS parsers."""
    # Remove BOM
    text = text.replace('\ufeff', '')
    # Remove zero-width spaces and joiners
    text = text.replace('\u200b', '')  # zero-width space
    text = text.replace('\u200c', '')  # zero-width non-joiner
    text = text.replace('\u200d', '')  # zero-width joiner
    text = text.replace('\u2060', '')  # word joiner
    text = text.replace('\ufffe', '')  # byte order mark variant
    # Remove soft hyphens
    text = text.replace('\u00ad', '')
    # Replace smart quotes with straight quotes (ATS-safe)
    text = text.replace('\u2018', "'").replace('\u2019', "'")  # single smart quotes
    text = text.replace('\u201c', '"').replace('\u201d', '"')  # double smart quotes
    # Replace en-dash and em-dash with ASCII hyphen (ATS-safe)
    text = text.replace('\u2013', '-')   # en-dash → hyphen
    text = text.replace('\u2014', '-')   # em-dash → hyphen
    # Replace ellipsis character with three dots
    text = text.replace('\u2026', '...')
    # Replace non-breaking space with regular space
    text = text.replace('\u00a0', ' ')
    # Remove any other control characters (except newline, tab)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Normalise multiple spaces to single
    text = re.sub(r' {2,}', ' ', text)
    return text


def fix_british_english(text: str) -> str:
    """Replace American English spellings with British English, context-aware."""
    for us, uk in US_TO_UK.items():
        pattern = re.compile(r'\b' + re.escape(us) + r'\b', re.IGNORECASE)
        def _replace(m, _uk=uk):
            orig = m.group()
            if orig[0].isupper():
                return _uk[0].upper() + _uk[1:]
            return _uk
        text = pattern.sub(_replace, text)

    # Context-aware "program" → "programme"
    # Only replace "program" when NOT preceded by tech context words
    def _replace_program(m):
        start = max(0, m.start() - 30)
        preceding = text[start:m.start()].lower()
        tech_words = ['computer', 'software', 'python', 'java', 'shell', 'c++',
                      'coding', 'code', 'script', 'executable', 'binary']
        for tw in tech_words:
            if tw in preceding:
                return m.group()  # keep "program" in tech context
        orig = m.group()
        replacement = 'programme'
        if orig[0].isupper():
            replacement = 'Programme'
        return replacement

    text = re.compile(r'\bprogram\b', re.IGNORECASE).sub(_replace_program, text)
    # "programming" should NEVER be changed
    # "programmer" should NEVER be changed
    # These are not in the dict so they're safe

    return text


def set_font(run, size=None, bold=False, italic=False, color=None):
    """Set font properties on a run. Sets both w:rFonts and fallback for ATS."""
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE_BODY
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or TEXT_COLOR
    # Set eastAsia and cs font names for full ATS compatibility
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)


def set_paragraph_spacing(para, before=0, after=0):
    """Set explicit spacing via XML — values in points, converted to twips.
    Line spacing is always 1.15x (276 twips, auto rule)."""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before * 20)))
    spacing.set(qn('w:after'), str(int(after * 20)))
    spacing.set(qn('w:line'), str(LINE_SPACING_VAL))
    spacing.set(qn('w:lineRule'), 'auto')


def add_horizontal_rule(doc):
    """Add a thin light-grey bottom border to the last paragraph (section divider)."""
    para = doc.paragraphs[-1]
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), DIVIDER_COLOR)
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_markdown(content: str) -> list[dict]:
    """Parse resume markdown into structured blocks."""
    blocks = []
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            blocks.append({"type": "blank"})
        elif stripped.startswith("#### "):
            blocks.append({"type": "subrole", "text": stripped[5:].strip()})
        elif stripped.startswith("### "):
            blocks.append({"type": "role", "text": stripped[4:].strip()})
        elif stripped.startswith("## "):
            blocks.append({"type": "section", "text": stripped[3:].strip()})
        elif stripped.startswith("# "):
            blocks.append({"type": "name", "text": stripped[2:].strip()})
        elif stripped.startswith("- ") or stripped.startswith("• "):
            blocks.append({"type": "bullet", "text": stripped[2:].strip()})
        else:
            blocks.append({"type": "body", "text": stripped})
    return blocks


def add_hyperlink(para, text, url, size=None, color=None):
    """Add a clickable hyperlink run to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    # Set font explicitly so it doesn't inherit default Hyperlink style weirdness
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rPr.append(rFonts)
    sz_el = OxmlElement('w:sz')
    sz_val = size or FONT_SIZE_BODY
    sz_el.set(qn('w:val'), str(int(sz_val.pt * 2)))  # half-points
    rPr.append(sz_el)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(sz_val.pt * 2)))
    rPr.append(szCs)
    link_color = color or HEADING_COLOR
    c_el = OxmlElement('w:color')
    c_el.set(qn('w:val'), str(link_color))
    rPr.append(c_el)
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


LINKEDIN_URL = ""         # set from --linkedin-url arg at runtime
CANDIDATE_AUTHOR = ""     # set from --candidate-name arg at runtime


def apply_inline_formatting(para, text: str, base_size=None, base_bold=False):
    """Apply **bold**, *italic* inline markdown, and LinkedIn hyperlink. No ghost empty runs."""
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            if inner:
                run = para.add_run(inner)
                set_font(run, size=base_size, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            inner = part[1:-1]
            if inner:
                run = para.add_run(inner)
                set_font(run, size=base_size, italic=True)
        else:
            if not part:
                continue
            # Check if this part contains the word "LinkedIn" to make it a hyperlink
            if "LinkedIn" in part:
                before, _, after = part.partition("LinkedIn")
                if before:
                    run = para.add_run(before)
                    set_font(run, size=base_size, bold=base_bold)
                add_hyperlink(para, "LinkedIn", LINKEDIN_URL, size=base_size)
                if after:
                    run = para.add_run(after)
                    set_font(run, size=base_size, bold=base_bold)
            else:
                run = para.add_run(part)
                set_font(run, size=base_size, bold=base_bold)


def create_bullet_numbering(doc):
    """Create a proper bullet list numbering definition in the document.
    Returns the abstractNumId to use."""
    numbering_part = doc.part.numbering_part
    # Get the numbering element
    numbering = numbering_part.numbering_definitions._numbering

    # Create abstract numbering with bullet
    abstract_num_id = 100  # Use high number to avoid conflicts
    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), str(abstract_num_id))

    # Multi-level definition (only level 0 needed)
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')

    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)

    numFmt = OxmlElement('w:numFmt')
    numFmt.set(qn('w:val'), 'bullet')
    lvl.append(numFmt)

    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), '\u2022')  # bullet character
    lvl.append(lvlText)

    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(qn('w:val'), 'left')
    lvl.append(lvlJc)

    # Paragraph properties for the bullet level
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(BULLET_INDENT))
    ind.set(qn('w:hanging'), str(BULLET_INDENT))
    pPr.append(ind)
    lvl.append(pPr)

    # Run properties for the bullet character
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Symbol')
    rFonts.set(qn('w:hAnsi'), 'Symbol')
    rFonts.set(qn('w:hint'), 'default')
    rPr.append(rFonts)
    lvl.append(rPr)

    abstractNum.append(lvl)
    # Insert before first num element
    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        numbering.insert(list(numbering).index(first_num), abstractNum)
    else:
        numbering.append(abstractNum)

    # Create num element referencing the abstract
    num_id = 100
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstractNumRef = OxmlElement('w:abstractNumId')
    abstractNumRef.set(qn('w:val'), str(abstract_num_id))
    num.append(abstractNumRef)
    numbering.append(num)

    return num_id


def add_bullet(doc, text, num_id=None, is_first=False):
    """Add a bullet point. First bullet in a group gets more space before."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if num_id is not None:
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        numIdEl = OxmlElement('w:numId')
        numIdEl.set(qn('w:val'), str(num_id))
        numPr.append(numIdEl)
        pPr.insert(0, numPr)
    else:
        text = '\u2022 ' + text

    apply_inline_formatting(p, text, base_size=FONT_SIZE_BODY)

    # Spacing: first bullet 12pt before, rest 1pt before, no after
    bef = 12 if is_first else 1
    # Set spacing with no after (0)
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(bef * 20)))
    # No w:after — let it inherit (matches your edit where after='-')
    spacing.set(qn('w:line'), str(LINE_SPACING_VAL))
    spacing.set(qn('w:lineRule'), 'auto')

    # Indent: 340 twips left
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), str(BULLET_INDENT))
    ind.set(qn('w:hanging'), str(BULLET_INDENT))

    return p


def setup_page(doc):
    """Configure A4 page with professional margins. Clear headers/footers."""
    section = doc.sections[0]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOT
    # Ensure no header/footer content (ATS can't read these)
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = True
    footer = section.footer
    footer.is_linked_to_previous = True
    # Remove the default empty paragraph that Document() creates
    # This phantom paragraph eats space at the top and causes the "clipped name" issue
    if doc.paragraphs and not doc.paragraphs[0].text:
        p_element = doc.paragraphs[0]._p
        p_element.getparent().remove(p_element)
    return section


def clean_metadata(doc):
    """Remove python-docx fingerprints from document metadata."""
    props = doc.core_properties
    author = CANDIDATE_AUTHOR or "Candidate"
    props.author = author
    props.last_modified_by = author
    props.title = ""
    props.subject = ""
    props.category = ""
    props.comments = ""
    props.keywords = ""
    # Set dates to now
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc)
    props.created = now
    props.modified = now


def setup_normal_style(doc):
    """Configure the Normal style as the base for all text."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY
    font.color.rgb = TEXT_COLOR
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(13)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Set rFonts on the style element for full ATS coverage
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)



def build_docx(content: str, output_path: str, apply_british_english: bool = True):
    """Build the resume DOCX from markdown content."""
    content = sanitise_text(content)
    if apply_british_english:
        content = fix_british_english(content)

    doc = Document()
    setup_page(doc)
    setup_normal_style(doc)

    num_id = None
    try:
        num_id = create_bullet_numbering(doc)
    except Exception:
        pass

    blocks = parse_markdown(content)
    prev_type = None
    in_header = True
    first_after_section = False
    is_first_bullet = True
    is_designation = False  # next body line after name is the designation

    for block in blocks:
        btype = block.get("type")
        text = block.get("text", "")

        if btype == "blank":
            if prev_type == "bullet":
                is_first_bullet = True
            prev_type = btype
            continue

        if btype == "section":
            in_header = False
            first_after_section = True
            is_first_bullet = True

        if btype == "name":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.upper())
            set_font(run, size=FONT_SIZE_NAME, bold=True, color=NAME_COLOR)
            set_paragraph_spacing(p, before=12, after=4)
            is_designation = True

        elif btype == "section":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text.upper())
            set_font(run, size=FONT_SIZE_SECTION, bold=True, color=HEADING_COLOR)
            set_paragraph_spacing(p, before=11, after=3)
            add_horizontal_rule(doc)

        elif btype == "role":
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font(run, size=FONT_SIZE_ROLE, bold=True, color=ROLE_COLOR)
            set_paragraph_spacing(p, before=7, after=1)
            is_first_bullet = True

        elif btype == "subrole":
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font(run, size=FONT_SIZE_BODY, italic=True)
            set_paragraph_spacing(p, before=0, after=3)
            is_first_bullet = True

        elif btype == "bullet":
            add_bullet(doc, text, num_id=num_id, is_first=is_first_bullet)
            is_first_bullet = False

        elif btype == "body":
            p = doc.add_paragraph()
            if in_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if is_designation:
                    apply_inline_formatting(p, text, base_size=FONT_SIZE_DESIGNATION)
                    is_designation = False
                else:
                    apply_inline_formatting(p, text, base_size=FONT_SIZE_CONTACT)
                set_paragraph_spacing(p, before=0, after=3)
            else:
                is_skills = '**' in text and len(text.split('**')) > 2 and ':' in text.split('**')[1]
                if is_skills:
                    _build_skills_line(p, text)
                    set_paragraph_spacing(p, before=12, after=3)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    apply_inline_formatting(p, text, base_size=FONT_SIZE_BODY)
                    set_paragraph_spacing(p, before=12 if first_after_section else 0, after=3)
                first_after_section = False

        prev_type = btype

    clean_metadata(doc)
    if os.path.dirname(output_path): os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print("Done: " + output_path)
    return output_path


def _build_skills_line(para, text):
    """Build a skills line: bold category + tab + values, with hanging indent."""
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = para._p.get_or_add_pPr()
    # Ensure a deterministic two-column skills layout:
    # - explicit tab stop where the value column starts
    # - hanging indent aligned to the same position for wrapped lines
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), str(SKILLS_INDENT))
    ind.set(qn('w:hanging'), str(SKILLS_INDENT))

    tabs = pPr.find(qn('w:tabs'))
    if tabs is not None:
        pPr.remove(tabs)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(SKILLS_INDENT))
    tabs.append(tab)
    pPr.append(tabs)

    m = re.match(r'\*\*(.+?)\*\*\s*(.*)', text)
    if m:
        category = m.group(1)
        values = m.group(2)
        run = para.add_run(category)
        set_font(run, size=FONT_SIZE_BODY, bold=True)
        run2 = para.add_run('\t')
        set_font(run2, size=FONT_SIZE_BODY)
        if values:
            run3 = para.add_run(values)
            set_font(run3, size=FONT_SIZE_BODY)
    else:
        apply_inline_formatting(para, text, base_size=FONT_SIZE_BODY)


def build_cover_letter_docx(content: str, output_path: str, apply_british_english: bool = True):
    """Build a cover letter DOCX with professional formatting."""
    content = sanitise_text(content)
    if apply_british_english:
        content = fix_british_english(content)

    doc = Document()
    setup_page(doc)
    setup_normal_style(doc)

    num_id = None
    try:
        num_id = create_bullet_numbering(doc)
    except Exception:
        pass

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=2)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:].strip().upper())
            set_font(run, size=FONT_SIZE_NAME, bold=True, color=NAME_COLOR)
            set_paragraph_spacing(p, before=0, after=4)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:].strip())
            set_font(run, size=FONT_SIZE_BODY)
            set_paragraph_spacing(p, before=0, after=3)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            add_bullet(doc, stripped[2:].strip(), num_id=num_id)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            apply_inline_formatting(p, stripped, base_size=FONT_SIZE_BODY)
            set_paragraph_spacing(p, before=0, after=6)

    clean_metadata(doc)
    if os.path.dirname(output_path): os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print("Done: " + output_path)
    return output_path


def convert_to_pdf(docx_path: str) -> str | None:
    """Convert DOCX to PDF using LibreOffice (if available) or docx2pdf."""
    pdf_path = docx_path.replace('.docx', '.pdf')
    out_dir = os.path.dirname(docx_path)

    for lo_bin in ['/Applications/LibreOffice.app/Contents/MacOS/soffice',
                   'soffice', 'libreoffice']:
        try:
            result = subprocess.run(
                [lo_bin, '--headless', '--convert-to', 'pdf',
                 '--outdir', out_dir, docx_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and os.path.exists(pdf_path):
                print(f"PDF saved: {pdf_path}")
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    try:
        import docx2pdf
        docx2pdf.convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            print(f"PDF saved: {pdf_path}")
            return pdf_path
    except (ImportError, Exception):
        pass

    print("PDF conversion not available. Install LibreOffice or docx2pdf.")
    return None


def main():
    parser = argparse.ArgumentParser(
        description="Generate ATS-safe .docx resume or cover letter"
    )
    parser.add_argument("--company", required=False, default="",
                        help="Company name (used for context)")
    parser.add_argument("--candidate-name", required=False, default="",
                        help="Candidate full name for document metadata and LinkedIn hyperlink label")
    parser.add_argument("--linkedin-url", required=False, default="",
                        help="Candidate LinkedIn profile URL for hyperlink")
    parser.add_argument("--market", required=False, default="uk", choices=["us", "uk"],
                        help="Target job market: 'us' (American English) or 'uk' (British English, default)")
    parser.add_argument("--content-file", help="Path to markdown content file")
    parser.add_argument("--inline-content", help="Content as inline string")
    parser.add_argument("--type", choices=["resume", "cover-letter"], default="resume",
                        help="Document type: 'resume' (default) or 'cover-letter'")
    parser.add_argument("--pdf", action="store_true",
                        help="Also generate a PDF version")
    args = parser.parse_args()

    global LINKEDIN_URL, CANDIDATE_AUTHOR
    LINKEDIN_URL = args.linkedin_url or ""
    CANDIDATE_AUTHOR = args.candidate_name or ""
    apply_british_english = (args.market == "uk")

    if args.content_file:
        with open(args.content_file, "r", encoding="utf-8") as f:
            content = f.read()
    elif args.inline_content:
        content = args.inline_content
    else:
        print("ERROR: Provide --content-file or --inline-content")
        sys.exit(1)

    company_slug = args.company.replace(" ", "_")

    if args.type == "cover-letter":
        output_path = "cover_letter.docx"
        docx_path = build_cover_letter_docx(content, output_path, apply_british_english)
    else:
        output_path = "resume.docx"
        docx_path = build_docx(content, output_path, apply_british_english)

    if args.pdf and docx_path:
        convert_to_pdf(docx_path)


if __name__ == "__main__":
    main()
