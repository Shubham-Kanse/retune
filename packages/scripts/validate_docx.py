#!/usr/bin/env python3
"""
validate_docx.py — Post-generation DOCX integrity validator.
Checks for every known ATS-killing issue in the generated document.

Usage:
    python3 scripts/validate_docx.py $WORKSPACE/resume.docx
"""

import sys
import os

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.")
    sys.exit(1)


def validate(path: str) -> list[dict]:
    """Validate a DOCX file. Returns list of issues found."""
    if not os.path.exists(path):
        return [{"severity": "CRITICAL", "check": "file_exists", "msg": f"File not found: {path}"}]

    doc = Document(path)
    issues = []

    def fail(severity, check, msg):
        issues.append({"severity": severity, "check": check, "msg": msg})

    # ── Page setup ──
    s = doc.sections[0]
    w_mm = round(s.page_width.mm)
    h_mm = round(s.page_height.mm)
    if w_mm != 210 or h_mm != 297:
        fail("CRITICAL", "page_size", f"Page is {w_mm}x{h_mm}mm, expected 210x297mm (A4)")

    for name, margin in [("left", s.left_margin), ("right", s.right_margin),
                         ("top", s.top_margin), ("bottom", s.bottom_margin)]:
        m_mm = round(margin.mm)
        if m_mm < 12 or m_mm > 26:
            fail("HIGH", "margins", f"{name} margin is {m_mm}mm, expected 12-25mm")

    # ── Metadata ──
    props = doc.core_properties
    if props.author and 'python' in props.author.lower():
        fail("HIGH", "metadata_author", f"Author contains 'python': '{props.author}'")
    if props.comments and ('generated' in props.comments.lower() or 'python' in props.comments.lower()):
        fail("HIGH", "metadata_comments", f"Comments reveal tool: '{props.comments}'")
    if props.last_modified_by and 'python' in props.last_modified_by.lower():
        fail("HIGH", "metadata_modifier", f"Last modified by reveals tool: '{props.last_modified_by}'")

    # ── Ghost characters ──
    GHOST_CHARS = {
        0x200B: "zero-width space",
        0x200C: "zero-width non-joiner",
        0x200D: "zero-width joiner",
        0x2060: "word joiner",
        0xFEFF: "BOM",
        0x00AD: "soft hyphen",
        0x2013: "en-dash",
        0x2014: "em-dash",
        0x2018: "left single smart quote",
        0x2019: "right single smart quote",
        0x201C: "left double smart quote",
        0x201D: "right double smart quote",
        0x00A0: "non-breaking space",
        0x2026: "ellipsis character",
    }

    for i, p in enumerate(doc.paragraphs):
        for j, ch in enumerate(p.text):
            code = ord(ch)
            if code in GHOST_CHARS:
                fail("CRITICAL", "ghost_char",
                     f"Para {i} pos {j}: U+{code:04X} ({GHOST_CHARS[code]}) in '{p.text[:50]}...'")
            elif code < 32 and ch not in '\n\r\t':
                fail("CRITICAL", "control_char",
                     f"Para {i} pos {j}: control char U+{code:04X}")

    # ── Empty runs ──
    for i, p in enumerate(doc.paragraphs):
        for j, r in enumerate(p.runs):
            if r.text == '':
                fail("HIGH", "empty_run",
                     f"Para {i} run {j}: empty text run (ghost run)")

    # ── Spacing inheritance ──
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue  # Skip blank separator paragraphs
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                fail("MEDIUM", "spacing_inherit",
                     f"Para {i}: no explicit spacing element")
        else:
            fail("MEDIUM", "spacing_inherit",
                 f"Para {i}: no pPr element at all")

    # ── Bullet numbering ──
    for i, p in enumerate(doc.paragraphs):
        pf = p.paragraph_format
        if pf.left_indent and pf.left_indent > 0 and p.text.strip():
            numPr = p._p.find('.//' + qn('w:numPr'))
            has_manual = p.text.startswith('\u2022') or p.text.startswith('•')
            has_tab = '\t' in p.text  # Skills lines use tab + hanging indent, not bullets
            if numPr is None and not has_manual and not has_tab:
                fail("MEDIUM", "bullet_numpr",
                     f"Para {i}: indented but no numPr or manual bullet")

    # ── Font consistency ──
    for i, p in enumerate(doc.paragraphs):
        for j, r in enumerate(p.runs):
            if r.font.name and r.font.name not in ('Arial', 'Calibri', 'Times New Roman', 'Georgia', 'Verdana'):
                fail("HIGH", "font_mismatch",
                     f"Para {i} run {j}: font '{r.font.name}' (expected ATS-safe font)")

    # ── Images/objects ──
    if len(doc.inline_shapes) > 0:
        fail("CRITICAL", "images", f"{len(doc.inline_shapes)} inline shapes (images) found")

    # ── Relationships check ──
    for rel_id, rel in doc.part.rels.items():
        rtype = rel.reltype.split('/')[-1]
        if rtype in ('image', 'chart', 'oleObject', 'hyperlink'):
            if rtype == 'hyperlink':
                pass  # Hyperlinks are OK
            else:
                fail("HIGH", "embedded_object", f"Embedded {rtype} found: {rel_id}")

    return issues


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/validate_docx.py <path-to-docx>")
        sys.exit(1)

    path = sys.argv[1]
    issues = validate(path)

    print(f"\n{'=' * 60}")
    print(f"DOCX INTEGRITY REPORT: {os.path.basename(path)}")
    print(f"{'=' * 60}")

    if not issues:
        print("\n✅ PERFECT — Zero issues found. Document is ATS-safe.\n")
        sys.exit(0)

    # Group by severity
    for severity in ["CRITICAL", "HIGH", "MEDIUM"]:
        sev_issues = [i for i in issues if i["severity"] == severity]
        if sev_issues:
            print(f"\n{'🚨' if severity == 'CRITICAL' else '⚠️' if severity == 'HIGH' else 'ℹ️'} {severity} ({len(sev_issues)} issues):")
            for issue in sev_issues:
                print(f"  [{issue['check']}] {issue['msg']}")

    critical = sum(1 for i in issues if i["severity"] == "CRITICAL")
    high = sum(1 for i in issues if i["severity"] == "HIGH")

    print(f"\nTotal: {len(issues)} issues ({critical} critical, {high} high)")

    if critical > 0:
        print("\n❌ FAIL — Critical issues must be fixed before submission.\n")
        sys.exit(2)
    elif high > 0:
        print("\n⚠️ WARNING — High issues should be reviewed.\n")
        sys.exit(1)
    else:
        print("\n✅ PASS — No critical or high issues.\n")
        sys.exit(0)


if __name__ == "__main__":
    main()
